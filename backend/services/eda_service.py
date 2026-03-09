"""
EDA Service — generates a Jupyter notebook, Word document, and cleaned CSV
from an uploaded dataset file.
"""

from __future__ import annotations

import logging
import os
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import matplotlib
matplotlib.use("Agg")  # non-interactive backend — must be set before importing pyplot

import matplotlib.pyplot as plt
import nbformat
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from nbformat.v4 import new_code_cell, new_markdown_cell, new_notebook
from scipy import stats

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

async def generate_eda(file_path: str, project_folder: str, job_id: str) -> Dict[str, Any]:
    """
    Run a full EDA pipeline for the uploaded file.

    Parameters
    ----------
    file_path      : absolute path to the uploaded dataset
    project_folder : base folder for the project (uploads/{user_id}/{project_id})
    job_id         : unique job identifier (used for output folder name)

    Returns
    -------
    dict with keys: output_folder, notebook_path, docx_path,
                    cleaned_csv_path, zip_path (may be None)
    """
    output_folder = os.path.join(project_folder, f"eda_{job_id}")
    os.makedirs(output_folder, exist_ok=True)

    # Copy raw input into output folder so the notebook can reference it via local relative path
    data_file_name = os.path.basename(file_path)
    dest_file = os.path.join(output_folder, data_file_name)
    if not os.path.exists(dest_file):
        shutil.copy2(file_path, dest_file)

    df = _read_file(file_path)
    stats_dict = _analyze_dataframe(df)
    findings = _generate_findings(df, stats_dict)

    notebook_path = os.path.join(output_folder, "eda_report.ipynb")
    _build_notebook(df, data_file_name, output_folder, notebook_path, stats_dict, findings)
    _execute_notebook(notebook_path)

    docx_path = os.path.join(output_folder, "eda_report.docx")
    _create_word_doc(df, stats_dict, findings, docx_path)

    cleaned_csv_path = os.path.join(output_folder, "cleaned_data.csv")
    cleaned_df = _clean_dataframe(df, stats_dict)
    cleaned_df.to_csv(cleaned_csv_path, index=False)

    # Save pipeline config as JSON
    pipeline_config_path = os.path.join(output_folder, "pipeline_config.json")
    _save_pipeline_config(stats_dict, cleaned_df, pipeline_config_path)

    # Always bundle key artifacts into a zip for easy download
    zip_path = _zip_artifacts(
        output_folder,
        {
            "eda_report.ipynb": notebook_path,
            "eda_report.docx": docx_path,
            "cleaned_data.csv": cleaned_csv_path,
            "pipeline_config.json": pipeline_config_path,
            os.path.basename(dest_file): dest_file,
        },
    )

    return {
        "output_folder": output_folder,
        "notebook_path": notebook_path,
        "docx_path": docx_path,
        "cleaned_csv_path": cleaned_csv_path,
        "zip_path": zip_path,
        "status": "completed",
    }


# ---------------------------------------------------------------------------
# File reading
# ---------------------------------------------------------------------------

def _detect_file_format(filepath: str):
    """Return the appropriate pandas read function based on file extension."""
    ext = Path(filepath).suffix.lower()
    mapping = {
        ".csv": lambda p: pd.read_csv(p, sep=None, engine="python"),
        ".tsv": lambda p: pd.read_csv(p, sep="\t"),
        ".xls": pd.read_excel,
        ".xlsx": pd.read_excel,
        ".json": pd.read_json,
        ".parquet": pd.read_parquet,
        ".data": lambda p: pd.read_csv(p, header=None, sep=None, engine="python"),
        ".file": lambda p: pd.read_csv(p, sep=None, engine="python"),
    }
    return mapping.get(ext, lambda p: pd.read_csv(p, sep=None, engine="python"))


def _read_file(filepath: str) -> pd.DataFrame:
    """Read a dataset file into a DataFrame, auto-detecting the format."""
    reader = _detect_file_format(filepath)
    df = reader(filepath)
    # Flatten multi-index columns (e.g., from some Excel files)
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = ["_".join(str(c) for c in col).strip() for col in df.columns]
    # Auto-name columns for headerless files (like .data)
    if all(isinstance(c, int) for c in df.columns):
        df.columns = [f"feature_{i}" if i < len(df.columns) - 1 else "target" for i in range(len(df.columns))]
    return df


# ---------------------------------------------------------------------------
# Analysis helpers
# ---------------------------------------------------------------------------

def _analyze_dataframe(df: pd.DataFrame) -> Dict[str, Any]:
    """Compute a comprehensive statistics dictionary for *df*."""
    rows, cols = df.shape
    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    cat_cols = df.select_dtypes(include=["object", "category"]).columns.tolist()

    # Missing values
    null_counts = df.isnull().sum().to_dict()
    null_pct = (df.isnull().mean() * 100).round(2).to_dict()

    # Descriptive stats
    desc = df.describe(include="all").to_dict()

    # Correlation matrix (numeric only)
    corr_matrix: Optional[pd.DataFrame] = None
    top_correlations: List[Tuple[str, str, float]] = []
    if len(num_cols) >= 2:
        corr_matrix = df[num_cols].corr()
        # Top 5 absolute correlations (excluding self-correlations)
        corr_pairs = (
            corr_matrix.abs()
            .unstack()
            .reset_index()
        )
        corr_pairs.columns = ["col1", "col2", "corr"]
        corr_pairs = corr_pairs[corr_pairs["col1"] < corr_pairs["col2"]]
        corr_pairs = corr_pairs.sort_values("corr", ascending=False)
        # Use signed values from original matrix for correct direction
        top_correlations = [
            (row["col1"], row["col2"], round(float(corr_matrix.loc[row["col1"], row["col2"]]), 4))
            for _, row in corr_pairs.head(5).iterrows()
        ]

    # Skewness
    skewness = {col: round(df[col].skew(), 4) for col in num_cols}

    # Outlier counts via IQR
    outlier_counts: Dict[str, int] = {}
    for col in num_cols:
        q1, q3 = df[col].quantile(0.25), df[col].quantile(0.75)
        iqr = q3 - q1
        mask = (df[col] < q1 - 1.5 * iqr) | (df[col] > q3 + 1.5 * iqr)
        outlier_counts[col] = int(mask.sum())

    # Date column detection
    date_cols: List[str] = []
    for col in df.columns:
        col_str = str(col).lower()
        if "date" in col_str or "time" in col_str or "year" in col_str:
            date_cols.append(col)

    # ---- Feature Engineering Analysis ----
    feature_engineering: Dict[str, Any] = {}

    # Interaction candidates: pairs with moderate-to-high correlation
    interaction_candidates: List[Tuple[str, str]] = []
    if len(num_cols) >= 2 and corr_matrix is not None:
        for c1, c2, corr_val in top_correlations:
            if 0.3 < abs(corr_val) < 0.95:
                interaction_candidates.append((c1, c2))
    feature_engineering["interaction_candidates"] = interaction_candidates[:5]

    # Date decomposition candidates
    date_decomp_cols: List[str] = []
    for col in df.columns:
        if pd.api.types.is_datetime64_any_dtype(df[col]):
            date_decomp_cols.append(col)
        elif df[col].dtype == object:
            sample = df[col].dropna().head(20)
            if len(sample) > 0:
                try:
                    pd.to_datetime(sample)
                    date_decomp_cols.append(col)
                except (ValueError, TypeError):
                    pass
    feature_engineering["date_decomp_cols"] = date_decomp_cols

    # Binning candidates: numeric columns with high cardinality
    binning_candidates: List[str] = []
    for col in num_cols:
        nunique = df[col].nunique()
        if nunique > 20 and not str(col).lower().endswith("id"):
            binning_candidates.append(col)
    feature_engineering["binning_candidates"] = binning_candidates[:8]

    # Log-transform candidates: positively skewed columns
    log_candidates: List[str] = []
    for col in num_cols:
        if col in skewness and skewness[col] > 1.0 and (df[col] > 0).all():
            log_candidates.append(col)
    feature_engineering["log_candidates"] = log_candidates

    # ---- Aggregation / GroupBy Analysis ----
    aggregation_analysis: Dict[str, Any] = {}

    # Identify good groupby columns (categorical with 2-20 unique values)
    groupby_cols: List[str] = []
    for col in cat_cols:
        nunique = df[col].nunique()
        if 2 <= nunique <= 20:
            groupby_cols.append(col)
    aggregation_analysis["groupby_cols"] = groupby_cols[:5]

    # Compute groupby summaries for top groupby cols
    groupby_summaries: Dict[str, Dict] = {}
    for grp_col in groupby_cols[:3]:
        if num_cols:
            agg_cols = num_cols[:5]
            try:
                summary = df.groupby(grp_col)[agg_cols].agg(["mean", "median", "std", "count"]).round(3)
                groupby_summaries[grp_col] = {
                    "agg_columns": agg_cols,
                    "n_groups": df[grp_col].nunique(),
                    "group_sizes": df[grp_col].value_counts().to_dict(),
                }
            except Exception:
                pass
    aggregation_analysis["groupby_summaries"] = groupby_summaries

    # ---- Trend Analysis ----
    trend_analysis: Dict[str, Any] = {}
    trend_cols_detected: List[str] = []
    for col in date_cols + date_decomp_cols:
        try:
            dt_series = pd.to_datetime(df[col], errors="coerce")
            valid_count = dt_series.notna().sum()
            if valid_count > 10:
                trend_cols_detected.append(col)
        except Exception:
            pass
    trend_analysis["time_columns"] = list(set(trend_cols_detected))

    # Detect monotonic trends in numeric columns
    monotonic_trends: Dict[str, str] = {}
    for col in num_cols[:10]:
        try:
            s = df[col].dropna()
            if len(s) >= 10:
                # Use Spearman correlation with index as proxy for trend
                from scipy.stats import spearmanr
                corr_coef, pval = spearmanr(range(len(s)), s)
                if pval < 0.05:
                    if corr_coef > 0.3:
                        monotonic_trends[col] = "increasing"
                    elif corr_coef < -0.3:
                        monotonic_trends[col] = "decreasing"
        except Exception:
            pass
    trend_analysis["monotonic_trends"] = monotonic_trends

    # ---- Pipeline Recommendations ----
    pipeline_recommendations: Dict[str, Any] = {}

    # Determine likely task type
    potential_target = None
    if cat_cols:
        # Last column is often the target
        last_col = df.columns[-1]
        if last_col in cat_cols and df[last_col].nunique() <= 20:
            potential_target = last_col
            pipeline_recommendations["suggested_task"] = "classification"
            pipeline_recommendations["suggested_target"] = last_col
            n_classes = df[last_col].nunique()
            pipeline_recommendations["n_classes"] = n_classes

            # Check class imbalance
            class_dist = df[last_col].value_counts(normalize=True)
            min_ratio = class_dist.min()
            pipeline_recommendations["class_imbalance"] = min_ratio < 0.1
            if n_classes == 2:
                pipeline_recommendations["suggested_models"] = ["RandomForest", "XGBoost", "LogisticRegression"]
                pipeline_recommendations["suggested_metric"] = "roc_auc" if min_ratio < 0.1 else "accuracy"
            else:
                pipeline_recommendations["suggested_models"] = ["RandomForest", "XGBoost", "GradientBoosting"]
                pipeline_recommendations["suggested_metric"] = "f1_weighted"
        elif last_col in num_cols:
            potential_target = last_col
            pipeline_recommendations["suggested_task"] = "regression"
            pipeline_recommendations["suggested_target"] = last_col
            pipeline_recommendations["suggested_models"] = ["RandomForestRegressor", "XGBoostRegressor", "GradientBoostingRegressor"]
            pipeline_recommendations["suggested_metric"] = "r2"
    elif num_cols:
        last_col = df.columns[-1]
        if last_col in num_cols:
            nunique = df[last_col].nunique()
            if nunique <= 15:
                pipeline_recommendations["suggested_task"] = "classification"
                pipeline_recommendations["suggested_target"] = last_col
                pipeline_recommendations["suggested_models"] = ["RandomForest", "XGBoost", "LogisticRegression"]
                pipeline_recommendations["suggested_metric"] = "f1_weighted"
            else:
                pipeline_recommendations["suggested_task"] = "regression"
                pipeline_recommendations["suggested_target"] = last_col
                pipeline_recommendations["suggested_models"] = ["RandomForestRegressor", "XGBoostRegressor", "GradientBoostingRegressor"]
                pipeline_recommendations["suggested_metric"] = "r2"

    if not pipeline_recommendations.get("suggested_task") and len(num_cols) >= 3:
        pipeline_recommendations["suggested_task"] = "clustering"
        pipeline_recommendations["suggested_models"] = ["KMeans", "DBSCAN", "GaussianMixture"]
        pipeline_recommendations["suggested_metric"] = "silhouette_score"

    # Imputer recommendation
    total_missing = sum(null_counts.values())
    missing_pct = (total_missing / (rows * cols)) * 100 if rows * cols > 0 else 0
    if missing_pct == 0:
        pipeline_recommendations["suggested_imputer"] = "none"
    elif missing_pct < 5:
        pipeline_recommendations["suggested_imputer"] = "SimpleImputer(strategy='median')"
    elif missing_pct < 15:
        pipeline_recommendations["suggested_imputer"] = "KNNImputer(n_neighbors=5)"
    else:
        pipeline_recommendations["suggested_imputer"] = "IterativeImputer(max_iter=10)"

    # Scaler recommendation
    highly_skewed_count = sum(1 for s in skewness.values() if abs(s) > 1)
    outlier_total = sum(outlier_counts.values())
    if outlier_total > rows * 0.05 * len(num_cols):
        pipeline_recommendations["suggested_scaler"] = "RobustScaler"
    elif highly_skewed_count > len(num_cols) * 0.5:
        pipeline_recommendations["suggested_scaler"] = "RobustScaler"
    else:
        pipeline_recommendations["suggested_scaler"] = "StandardScaler"

    # Encoder recommendation
    high_card_cats = [c for c in cat_cols if df[c].nunique() > 15]
    if high_card_cats:
        pipeline_recommendations["suggested_encoder"] = "OrdinalEncoder (high cardinality detected)"
    elif cat_cols:
        pipeline_recommendations["suggested_encoder"] = "OneHotEncoder"
    else:
        pipeline_recommendations["suggested_encoder"] = "none"

    # Transformers list
    transformers: List[str] = []
    if pipeline_recommendations.get("suggested_scaler"):
        transformers.append(pipeline_recommendations["suggested_scaler"])
    if cat_cols:
        transformers.append(pipeline_recommendations.get("suggested_encoder", "OneHotEncoder"))
    if pipeline_recommendations.get("suggested_imputer") != "none":
        transformers.append("SimpleImputer")
    pipeline_recommendations["suggested_transformers"] = transformers

    return {
        "shape": df.shape,
        "num_cols": num_cols,
        "cat_cols": cat_cols,
        "null_counts": null_counts,
        "null_pct": null_pct,
        "desc": desc,
        "corr_matrix": corr_matrix,
        "top_correlations": top_correlations,
        "skewness": skewness,
        "outlier_counts": outlier_counts,
        "date_cols": date_cols,
        "duplicates": int(df.duplicated().sum()),
        "feature_engineering": feature_engineering,
        "aggregation_analysis": aggregation_analysis,
        "trend_analysis": trend_analysis,
        "pipeline_recommendations": pipeline_recommendations,
    }


def _generate_findings(df: pd.DataFrame, stats: Dict[str, Any]) -> List[str]:
    """Generate human-readable finding strings from the analysis stats."""
    findings: List[str] = []
    rows, cols = stats["shape"]
    findings.append(f"Dataset contains {rows:,} rows and {cols} columns.")
    findings.append(
        f"{len(stats['num_cols'])} numeric column(s) and "
        f"{len(stats['cat_cols'])} categorical column(s) detected."
    )

    if stats["duplicates"] > 0:
        findings.append(f"Found {stats['duplicates']:,} duplicate rows ({stats['duplicates']/rows*100:.1f}%).")

    # High-null columns
    high_null = {col: pct for col, pct in stats["null_pct"].items() if pct > 20}
    if high_null:
        cols_list = ", ".join(f"{c} ({p}%)" for c, p in list(high_null.items())[:5])
        findings.append(f"High missing-value columns (>20%): {cols_list}.")

    # Top correlations
    for c1, c2, corr in stats["top_correlations"][:3]:
        direction = "positive" if corr > 0 else "negative"
        strength = "strong" if abs(corr) > 0.7 else "moderate"
        findings.append(f"{strength.capitalize()} {direction} correlation ({corr}) between '{c1}' and '{c2}'.")

    # Skewed features
    highly_skewed = {c: s for c, s in stats["skewness"].items() if abs(s) > 1}
    if highly_skewed:
        cols_list = ", ".join(f"{c} (skew={s})" for c, s in list(highly_skewed.items())[:5])
        findings.append(f"Highly skewed features: {cols_list}.")

    # Outliers
    high_outlier = {c: n for c, n in stats["outlier_counts"].items() if n > 0}
    if high_outlier:
        summary = ", ".join(f"{c}: {n}" for c, n in list(high_outlier.items())[:5])
        findings.append(f"Outliers detected (IQR method): {summary}.")

    # Date/time columns
    if stats["date_cols"]:
        findings.append(f"Potential time-series column(s) detected: {', '.join(stats['date_cols'])}.")

    # Feature engineering findings
    fe = stats.get("feature_engineering", {})
    if fe.get("interaction_candidates"):
        pairs = ", ".join(f"{c1}×{c2}" for c1, c2 in fe["interaction_candidates"][:3])
        findings.append(f"Suggested interaction features: {pairs}.")
    if fe.get("log_candidates"):
        findings.append(f"Log-transform candidates (positively skewed, all positive): {', '.join(fe['log_candidates'][:5])}.")
    if fe.get("date_decomp_cols"):
        findings.append(f"Date columns suitable for decomposition (year/month/day/dow): {', '.join(fe['date_decomp_cols'])}.")
    if fe.get("binning_candidates"):
        findings.append(f"Binning candidates (high-cardinality numeric): {', '.join(fe['binning_candidates'][:5])}.")

    # Trend analysis findings
    trend = stats.get("trend_analysis", {})
    if trend.get("monotonic_trends"):
        trend_list = ", ".join(f"{c} ({d})" for c, d in list(trend["monotonic_trends"].items())[:5])
        findings.append(f"Monotonic trends detected: {trend_list}.")

    # Aggregation findings
    agg = stats.get("aggregation_analysis", {})
    if agg.get("groupby_cols"):
        findings.append(f"Suitable groupby columns for aggregation: {', '.join(agg['groupby_cols'])}.")

    # Pipeline recommendation findings
    pr = stats.get("pipeline_recommendations", {})
    if pr.get("suggested_task"):
        findings.append(
            f"Recommended pipeline: task={pr['suggested_task']}, "
            f"models={pr.get('suggested_models', [])}, "
            f"metric={pr.get('suggested_metric', 'N/A')}, "
            f"imputer={pr.get('suggested_imputer', 'N/A')}, "
            f"scaler={pr.get('suggested_scaler', 'N/A')}."
        )
    if pr.get("class_imbalance"):
        findings.append("⚠ Class imbalance detected — consider SMOTE or stratified sampling.")

    return findings


# ---------------------------------------------------------------------------
# Notebook generation
# ---------------------------------------------------------------------------

def _build_notebook(
    df: pd.DataFrame,
    data_file_name: str,
    output_folder: str,
    notebook_path: str,
    stats: Dict[str, Any],
    findings: List[str],
) -> None:
    """Build and write a comprehensive EDA notebook to *notebook_path*."""
    cells = []
    fname = data_file_name

    # ---- Section 1: Setup & Data Loading ----
    cells.append(new_markdown_cell("# 📊 EDA Report\n## Section 1 — Setup & Data Loading"))
    cells.append(new_code_cell(
        "import warnings\n"
        "warnings.filterwarnings('ignore')\n"
        "import matplotlib\n"
        "matplotlib.use('Agg')\n"
        "import matplotlib.pyplot as plt\n"
        "import seaborn as sns\n"
        "import pandas as pd\n"
        "import numpy as np\n"
        "from scipy import stats\n"
        "from sklearn.decomposition import PCA\n"
        "from sklearn.preprocessing import StandardScaler\n"
        "from pathlib import Path\n"
        "import os\n\n"
        "sns.set_theme(style='whitegrid')\n"
        "%matplotlib inline\n"
        "data_file = Path(" + repr(data_file_name) + ")\n"
        "if data_file.suffix.lower() == '.csv':\n"
        "    df = pd.read_csv(data_file)\n"
        "elif data_file.suffix.lower() == '.tsv':\n"
        "    df = pd.read_csv(data_file, sep='\\t')\n"
        "elif data_file.suffix.lower() in ['.xls', '.xlsx']:\n"
        "    df = pd.read_excel(data_file)\n"
        "elif data_file.suffix.lower() == '.json':\n"
        "    df = pd.read_json(data_file)\n"
        "else:\n"
        "    raise ValueError('Unsupported file format')\n"
        "print('Shape:', df.shape)\n"
        "print('\\nData Types:\\n', df.dtypes)\n"
        "df.head()"
    ))

    # ---- Section 2: Basic Statistics ----
    cells.append(new_markdown_cell("## Section 2 — Basic Statistics"))
    cells.append(new_code_cell(
        "print('--- Descriptive Statistics ---')\n"
        "display(df.describe(include='all'))\n\n"
        "print('\\n--- Null Counts ---')\n"
        "null_info = pd.DataFrame({'count': df.isnull().sum(), 'pct': (df.isnull().mean()*100).round(2)})\n"
        "display(null_info[null_info['count'] > 0].sort_values('pct', ascending=False))\n\n"
        "print(f'\\nDuplicate rows: {df.duplicated().sum()}')"
    ))

    # ---- Section 3: Data Quality ----
    cells.append(new_markdown_cell("## Section 3 — Data Quality"))
    cells.append(new_code_cell(
        "import matplotlib.pyplot as plt\n"
        "import seaborn as sns\n"
        "import numpy as np\n\n"
        "# Missing value heatmap\n"
        "null_cols = df.columns[df.isnull().any()].tolist()\n"
        "if null_cols:\n"
        "    fig, ax = plt.subplots(figsize=(min(14, len(null_cols)+2), 5))\n"
        "    sns.heatmap(df[null_cols].isnull(), cbar=True, yticklabels=False, ax=ax, cmap='viridis')\n"
        "    ax.set_title('Missing Value Heatmap')\n"
        "    plt.tight_layout()\n"
        "    plt.savefig(os.path.join(r'" + output_folder + "', 'missing_heatmap.png'), dpi=100, bbox_inches='tight')\n"
        "    plt.show()\n"
        "    plt.close()\n"
        "else:\n"
        "    print('No missing values — dataset is complete!')\n\n"
        "# IQR outlier detection\n"
        "num_cols = df.select_dtypes(include=np.number).columns.tolist()\n"
        "outlier_summary = {}\n"
        "for col in num_cols:\n"
        "    q1, q3 = df[col].quantile(0.25), df[col].quantile(0.75)\n"
        "    iqr = q3 - q1\n"
        "    mask = (df[col] < q1 - 1.5*iqr) | (df[col] > q3 + 1.5*iqr)\n"
        "    outlier_summary[col] = mask.sum()\n"
        "outlier_df = pd.Series(outlier_summary, name='outlier_count').to_frame()\n"
        "print('\\nOutlier counts (IQR method):')\n"
        "display(outlier_df[outlier_df['outlier_count'] > 0])"
    ))

    # ---- Section 4: Univariate Analysis ----
    cells.append(new_markdown_cell("## Section 4 — Univariate Analysis"))
    cells.append(new_code_cell(
        "num_cols = df.select_dtypes(include=np.number).columns.tolist()\n"
        "cat_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()\n\n"
        "# Histograms for numeric columns\n"
        "if num_cols:\n"
        "    n = len(num_cols)\n"
        "    ncols = 3\n"
        "    nrows = (n + ncols - 1) // ncols\n"
        "    fig, axes = plt.subplots(nrows, ncols, figsize=(ncols*5, nrows*4))\n"
        "    axes = np.array(axes).flatten()\n"
        "    for i, col in enumerate(num_cols):\n"
        "        axes[i].hist(df[col].dropna(), bins=30, color='steelblue', edgecolor='white', alpha=0.8)\n"
        "        axes[i].set_title(col, fontsize=11)\n"
        "        axes[i].set_xlabel('')\n"
        "    for j in range(i+1, len(axes)):\n"
        "        axes[j].set_visible(False)\n"
        "    plt.suptitle('Numeric Feature Distributions', fontsize=14, y=1.01)\n"
        "    plt.tight_layout()\n"
        "    plt.savefig(os.path.join(r'" + output_folder + "', 'numeric_histograms.png'), dpi=100, bbox_inches='tight')\n"
        "    plt.show(); plt.close()\n\n"
        "# Bar charts for categorical columns\n"
        "if cat_cols:\n"
        "    n = min(len(cat_cols), 9)\n"
        "    ncols = 3\n"
        "    nrows = (n + ncols - 1) // ncols\n"
        "    fig, axes = plt.subplots(nrows, ncols, figsize=(ncols*5, nrows*4))\n"
        "    axes = np.array(axes).flatten()\n"
        "    for i, col in enumerate(cat_cols[:9]):\n"
        "        vc = df[col].value_counts().head(15)\n"
        "        axes[i].barh(vc.index.astype(str), vc.values, color='coral', edgecolor='white')\n"
        "        axes[i].set_title(col, fontsize=11)\n"
        "    for j in range(i+1, len(axes)):\n"
        "        axes[j].set_visible(False)\n"
        "    plt.suptitle('Categorical Feature Value Counts', fontsize=14, y=1.01)\n"
        "    plt.tight_layout()\n"
        "    plt.savefig(os.path.join(r'" + output_folder + "', 'categorical_bars.png'), dpi=100, bbox_inches='tight')\n"
        "    plt.show(); plt.close()"
    ))

    # ---- Section 5: Bivariate Analysis ----
    cells.append(new_markdown_cell("## Section 5 — Bivariate Analysis"))
    cells.append(new_code_cell(
        "num_cols = df.select_dtypes(include=np.number).columns.tolist()\n\n"
        "# Correlation heatmap\n"
        "if len(num_cols) >= 2:\n"
        "    corr = df[num_cols].corr()\n"
        "    mask = np.triu(np.ones_like(corr, dtype=bool))\n"
        "    fig, ax = plt.subplots(figsize=(max(8, len(num_cols)), max(6, len(num_cols)-1)))\n"
        "    sns.heatmap(corr, mask=mask, annot=True, fmt='.2f', cmap='coolwarm',\n"
        "                center=0, square=True, linewidths=0.5, ax=ax)\n"
        "    ax.set_title('Correlation Heatmap', fontsize=14)\n"
        "    plt.tight_layout()\n"
        "    plt.savefig(os.path.join(r'" + output_folder + "', 'correlation_heatmap.png'), dpi=100, bbox_inches='tight')\n"
        "    plt.show(); plt.close()\n\n"
        "    # Scatter plots for top correlated pairs\n"
        "    corr_pairs = (corr.abs().unstack().reset_index())\n"
        "    corr_pairs.columns = ['col1','col2','corr']\n"
        "    corr_pairs = corr_pairs[corr_pairs['col1'] < corr_pairs['col2']].sort_values('corr', ascending=False)\n"
        "    top_pairs = corr_pairs.head(min(6, len(corr_pairs)))\n"
        "    if not top_pairs.empty:\n"
        "        n = len(top_pairs)\n"
        "        ncols = min(3, n)\n"
        "        nrows = (n + ncols - 1) // ncols\n"
        "        fig, axes = plt.subplots(nrows, ncols, figsize=(ncols*5, nrows*4))\n"
        "        axes = np.array(axes).flatten()\n"
        "        for i, (_, row) in enumerate(top_pairs.iterrows()):\n"
        "            axes[i].scatter(df[row['col1']], df[row['col2']], alpha=0.4, s=15, color='steelblue')\n"
        "            axes[i].set_xlabel(row['col1']); axes[i].set_ylabel(row['col2'])\n"
        "            axes[i].set_title(f\"r={row['corr']:.2f}\", fontsize=10)\n"
        "        for j in range(i+1, len(axes)):\n"
        "            axes[j].set_visible(False)\n"
        "        plt.suptitle('Top Correlated Pairs', fontsize=13)\n"
        "        plt.tight_layout()\n"
        "        plt.savefig(os.path.join(r'" + output_folder + "', 'scatter_pairs.png'), dpi=100, bbox_inches='tight')\n"
        "        plt.show(); plt.close()\n\n"
        "# Box plots\n"
        "if len(num_cols) >= 1:\n"
        "    ncols = 3\n"
        "    n = min(len(num_cols), 9)\n"
        "    nrows = (n + ncols - 1) // ncols\n"
        "    fig, axes = plt.subplots(nrows, ncols, figsize=(ncols*4, nrows*3))\n"
        "    axes = np.array(axes).flatten()\n"
        "    for i, col in enumerate(num_cols[:9]):\n"
        "        axes[i].boxplot(df[col].dropna(), patch_artist=True,\n"
        "                        boxprops=dict(facecolor='lightblue'))\n"
        "        axes[i].set_title(col, fontsize=10)\n"
        "    for j in range(i+1, len(axes)):\n"
        "        axes[j].set_visible(False)\n"
        "    plt.suptitle('Box Plots (Outlier View)', fontsize=13)\n"
        "    plt.tight_layout()\n"
        "    plt.savefig(os.path.join(r'" + output_folder + "', 'boxplots.png'), dpi=100, bbox_inches='tight')\n"
        "    plt.show(); plt.close()"
    ))

    # ---- Section 6: Multivariate Analysis ----
    cells.append(new_markdown_cell("## Section 6 — Multivariate Analysis"))
    cells.append(new_code_cell(
        "from sklearn.decomposition import PCA\n"
        "from sklearn.preprocessing import StandardScaler\n\n"
        "num_cols = df.select_dtypes(include=np.number).columns.tolist()\n\n"
        "# Pair plot for top-5 numeric features\n"
        "if len(num_cols) >= 3:\n"
        "    top5 = num_cols[:5]\n"
        "    pair_df = df[top5].dropna()\n"
        "    if len(pair_df) > 2000:\n"
        "        pair_df = pair_df.sample(2000, random_state=42)\n"
        "    pair_fig = sns.pairplot(pair_df, diag_kind='kde', plot_kws={'alpha': 0.3, 's': 10})\n"
        "    pair_fig.fig.suptitle('Pair Plot — Top Numeric Features', y=1.01, fontsize=13)\n"
        "    pair_fig.savefig(os.path.join(r'" + output_folder + "', 'pairplot.png'), dpi=80, bbox_inches='tight')\n"
        "    plt.show(); plt.close()\n\n"
        "# PCA for datasets with > 5 numeric columns\n"
        "if len(num_cols) > 5:\n"
        "    pca_df = df[num_cols].dropna()\n"
        "    scaler = StandardScaler()\n"
        "    X_scaled = scaler.fit_transform(pca_df)\n"
        "    pca = PCA(n_components=2, random_state=42)\n"
        "    components = pca.fit_transform(X_scaled)\n"
        "    fig, ax = plt.subplots(figsize=(8, 6))\n"
        "    ax.scatter(components[:, 0], components[:, 1], alpha=0.4, s=15, c='steelblue')\n"
        "    ax.set_xlabel(f'PC1 ({pca.explained_variance_ratio_[0]*100:.1f}% variance)')\n"
        "    ax.set_ylabel(f'PC2 ({pca.explained_variance_ratio_[1]*100:.1f}% variance)')\n"
        "    ax.set_title('PCA — First Two Principal Components')\n"
        "    plt.tight_layout()\n"
        "    plt.savefig(os.path.join(r'" + output_folder + "', 'pca_plot.png'), dpi=100, bbox_inches='tight')\n"
        "    plt.show(); plt.close()\n"
        "    print(f'PCA explained variance: {pca.explained_variance_ratio_}')"
    ))

    # ---- Section 7: Time Series Detection ----
    cells.append(new_markdown_cell("## Section 7 — Time Series Detection"))
    cells.append(new_code_cell(
        "import re\n"
        "date_candidates = [c for c in df.columns if re.search(r'date|time|year|month|day', c, re.I)]\n"
        "if date_candidates:\n"
        "    for dcol in date_candidates:\n"
        "        try:\n"
        "            df[dcol] = pd.to_datetime(df[dcol], errors='coerce')\n"
        "            ts = df.dropna(subset=[dcol]).set_index(dcol).sort_index()\n"
        "            num_cols = ts.select_dtypes(include=np.number).columns.tolist()\n"
        "            if num_cols:\n"
        "                n = min(len(num_cols), 3)\n"
        "                fig, axes = plt.subplots(n, 1, figsize=(12, 4*n))\n"
        "                if n == 1: axes = [axes]\n"
        "                for i, col in enumerate(num_cols[:n]):\n"
        "                    axes[i].plot(ts.index, ts[col], linewidth=0.8, color='steelblue')\n"
        "                    axes[i].set_title(f'{col} over time')\n"
        "                    axes[i].set_xlabel(dcol)\n"
        "                plt.suptitle(f'Time Series Trends (index: {dcol})', fontsize=13)\n"
        "                plt.tight_layout()\n"
        "                plt.savefig(os.path.join(r'" + output_folder + "', 'time_series.png'), dpi=100, bbox_inches='tight')\n"
        "                plt.show(); plt.close()\n"
        "        except Exception as e:\n"
        "            print(f'Could not parse {dcol} as datetime: {e}')\n"
        "else:\n"
        "    print('No date/time columns detected — skipping time series analysis.')"
    ))

    # ---- Section 8: Key Findings ----
    findings_md = "\n".join(f"- {f}" for f in findings)
    cells.append(new_markdown_cell(
        f"## Section 8 — Key Findings\n\n{findings_md}"
    ))
    cells.append(new_code_cell(
        "# Summary statistics overview\n"
        "print('=== Dataset Summary ===')\n"
        "print(f'Rows: {df.shape[0]:,}  |  Columns: {df.shape[1]}')\n"
        "print(f'Numeric columns : {len(df.select_dtypes(include=np.number).columns)}')\n"
        "print(f'Categorical cols: {len(df.select_dtypes(include=[\"object\",\"category\"]).columns)}')\n"
        "print(f'Total missing   : {df.isnull().sum().sum():,}')\n"
        "print(f'Duplicate rows  : {df.duplicated().sum():,}')"
    ))

    # ---- Section 9: Data Cleaning ----
    cells.append(new_markdown_cell("## Section 9 — Data Cleaning"))
    cells.append(new_code_cell(
        "cleaned = df.copy()\n\n"
        "# Drop duplicates\n"
        "before = len(cleaned)\n"
        "cleaned = cleaned.drop_duplicates()\n"
        "print(f'Removed {before - len(cleaned)} duplicate rows.')\n\n"
        "# Fill numeric NaN with median\n"
        "num_cols = cleaned.select_dtypes(include=np.number).columns.tolist()\n"
        "for col in num_cols:\n"
        "    if cleaned[col].isnull().any():\n"
        "        med = cleaned[col].median()\n"
        "        cleaned[col] = cleaned[col].fillna(med)\n\n"
        "# Fill categorical NaN with mode\n"
        "cat_cols = cleaned.select_dtypes(include=['object','category']).columns.tolist()\n"
        "for col in cat_cols:\n"
        "    if cleaned[col].isnull().any():\n"
        "        mode_val = cleaned[col].mode()\n"
        "        if not mode_val.empty:\n"
        "            cleaned[col] = cleaned[col].fillna(mode_val[0])\n\n"
        "# Remove IQR outliers from numeric columns (cap to 1.5*IQR)\n"
        "for col in num_cols:\n"
        "    q1, q3 = cleaned[col].quantile(0.25), cleaned[col].quantile(0.75)\n"
        "    iqr = q3 - q1\n"
        "    lower, upper = q1 - 1.5*iqr, q3 + 1.5*iqr\n"
        "    cleaned[col] = cleaned[col].clip(lower, upper)\n\n"
        "print(f'Cleaned shape: {cleaned.shape}')\n"
        "print(f'Remaining nulls: {cleaned.isnull().sum().sum()}')\n"
        "cleaned.head()"
    ))

    # ---- Section 10: Feature Engineering ----
    fe = stats.get("feature_engineering", {})
    cells.append(new_markdown_cell("## Section 10 — Feature Engineering"))
    fe_code = (
        "# --- Feature Engineering ---\n"
        "import warnings\n"
        "warnings.filterwarnings('ignore')\n"
        "fe_df = cleaned.copy()\n"
        "num_cols = fe_df.select_dtypes(include=np.number).columns.tolist()\n"
        "cat_cols = fe_df.select_dtypes(include=['object','category']).columns.tolist()\n"
        "new_features = []\n\n"
    )

    # Interaction features
    interaction_candidates = fe.get("interaction_candidates", [])
    if interaction_candidates:
        fe_code += "# Interaction features (product of correlated pairs)\n"
        for c1, c2 in interaction_candidates[:5]:
            safe_name = f"{c1}_x_{c2}"
            fe_code += f"if '{c1}' in fe_df.columns and '{c2}' in fe_df.columns:\n"
            fe_code += f"    fe_df['{safe_name}'] = fe_df['{c1}'] * fe_df['{c2}']\n"
            fe_code += f"    new_features.append('{safe_name}')\n"
        fe_code += "\n"

    # Log transform candidates
    log_candidates = fe.get("log_candidates", [])
    if log_candidates:
        fe_code += "# Log-transform for skewed positive features\n"
        for col in log_candidates[:5]:
            fe_code += f"if '{col}' in fe_df.columns and (fe_df['{col}'] > 0).all():\n"
            fe_code += f"    fe_df['{col}_log'] = np.log1p(fe_df['{col}'])\n"
            fe_code += f"    new_features.append('{col}_log')\n"
        fe_code += "\n"

    # Date decomposition
    date_decomp_cols = fe.get("date_decomp_cols", [])
    if date_decomp_cols:
        fe_code += "# Date decomposition\n"
        for col in date_decomp_cols[:3]:
            fe_code += f"try:\n"
            fe_code += f"    fe_df['{col}'] = pd.to_datetime(fe_df['{col}'], errors='coerce')\n"
            fe_code += f"    fe_df['{col}_year'] = fe_df['{col}'].dt.year\n"
            fe_code += f"    fe_df['{col}_month'] = fe_df['{col}'].dt.month\n"
            fe_code += f"    fe_df['{col}_dayofweek'] = fe_df['{col}'].dt.dayofweek\n"
            fe_code += f"    fe_df['{col}_quarter'] = fe_df['{col}'].dt.quarter\n"
            fe_code += f"    new_features.extend(['{col}_year', '{col}_month', '{col}_dayofweek', '{col}_quarter'])\n"
            fe_code += f"except Exception:\n"
            fe_code += f"    pass\n"
        fe_code += "\n"

    # Binning candidates
    binning_candidates = fe.get("binning_candidates", [])
    if binning_candidates:
        fe_code += "# Quantile binning for high-cardinality numeric features\n"
        for col in binning_candidates[:5]:
            fe_code += f"if '{col}' in fe_df.columns:\n"
            fe_code += f"    try:\n"
            fe_code += f"        fe_df['{col}_binned'] = pd.qcut(fe_df['{col}'], q=5, labels=['very_low','low','mid','high','very_high'], duplicates='drop')\n"
            fe_code += f"        new_features.append('{col}_binned')\n"
            fe_code += f"    except Exception:\n"
            fe_code += f"        pass\n"
        fe_code += "\n"

    fe_code += (
        "print(f'New features created: {len(new_features)}')\n"
        "if new_features:\n"
        "    print('Features:', new_features)\n"
        "    display(fe_df[new_features].describe())\n"
        "cleaned = fe_df  # Update cleaned with engineered features\n"
    )
    cells.append(new_code_cell(fe_code))

    # ---- Section 11: Trend Analysis & Aggregations ----
    cells.append(new_markdown_cell("## Section 11 — Trend Analysis & Aggregations"))

    # Aggregation / GroupBy code
    agg = stats.get("aggregation_analysis", {})
    groupby_cols = agg.get("groupby_cols", [])
    trend = stats.get("trend_analysis", {})

    agg_code = (
        "# --- Aggregation & GroupBy Analysis ---\n"
        "num_cols = cleaned.select_dtypes(include=np.number).columns.tolist()\n"
        "cat_cols = cleaned.select_dtypes(include=['object','category']).columns.tolist()\n\n"
    )

    if groupby_cols:
        for grp_col in groupby_cols[:3]:
            agg_code += f"# GroupBy: {grp_col}\n"
            agg_code += f"if '{grp_col}' in cleaned.columns:\n"
            agg_code += f"    agg_cols = [c for c in num_cols if c != '{grp_col}'][:5]\n"
            agg_code += f"    if agg_cols:\n"
            agg_code += f"        grp_stats = cleaned.groupby('{grp_col}')[agg_cols].agg(['mean', 'median', 'std', 'min', 'max'])\n"
            agg_code += f"        print(f'\\n=== Aggregation by {grp_col} ===')\n"
            agg_code += f"        display(grp_stats.round(3))\n\n"
            agg_code += f"        # Bar plot of mean values per group\n"
            agg_code += f"        plot_cols = agg_cols[:3]\n"
            agg_code += f"        means = cleaned.groupby('{grp_col}')[plot_cols].mean()\n"
            agg_code += f"        fig, axes = plt.subplots(1, len(plot_cols), figsize=(5*len(plot_cols), 4))\n"
            agg_code += f"        if len(plot_cols) == 1: axes = [axes]\n"
            agg_code += f"        for i, col in enumerate(plot_cols):\n"
            agg_code += f"            means[col].plot(kind='bar', ax=axes[i], color='steelblue', edgecolor='white')\n"
            agg_code += f"            axes[i].set_title(f'Mean {{col}} by {grp_col}', fontsize=11)\n"
            agg_code += f"            axes[i].set_xlabel('{grp_col}')\n"
            agg_code += f"            axes[i].tick_params(axis='x', rotation=45)\n"
            agg_code += f"        plt.tight_layout()\n"
            agg_code += f"        plt.savefig(os.path.join(r'{output_folder}', 'groupby_{grp_col}.png'), dpi=100, bbox_inches='tight')\n"
            agg_code += f"        plt.show(); plt.close()\n\n"
    else:
        agg_code += "# No suitable categorical columns for groupby (need 2-20 unique values)\n"
        agg_code += "print('No categorical columns suitable for groupby aggregation.')\n\n"

    # Trend analysis
    agg_code += "# --- Trend Analysis ---\n"
    time_cols = trend.get("time_columns", [])
    monotonic = trend.get("monotonic_trends", {})

    if time_cols:
        for tcol in time_cols[:2]:
            agg_code += f"try:\n"
            agg_code += f"    ts_col = pd.to_datetime(cleaned['{tcol}'], errors='coerce')\n"
            agg_code += f"    valid_mask = ts_col.notna()\n"
            agg_code += f"    if valid_mask.sum() > 10:\n"
            agg_code += f"        ts_df = cleaned[valid_mask].copy()\n"
            agg_code += f"        ts_df['{tcol}'] = ts_col[valid_mask]\n"
            agg_code += f"        ts_df = ts_df.sort_values('{tcol}')\n"
            agg_code += f"        ts_num = ts_df.select_dtypes(include=np.number).columns.tolist()[:3]\n"
            agg_code += f"        if ts_num:\n"
            agg_code += f"            fig, axes = plt.subplots(len(ts_num), 1, figsize=(12, 4*len(ts_num)))\n"
            agg_code += f"            if len(ts_num) == 1: axes = [axes]\n"
            agg_code += f"            for i, col in enumerate(ts_num):\n"
            agg_code += f"                axes[i].plot(ts_df['{tcol}'], ts_df[col], linewidth=0.8, alpha=0.6, label='Raw')\n"
            agg_code += f"                # Rolling mean for trend\n"
            agg_code += f"                window = max(5, len(ts_df) // 20)\n"
            agg_code += f"                rolling = ts_df[col].rolling(window=window, center=True).mean()\n"
            agg_code += f"                axes[i].plot(ts_df['{tcol}'], rolling, linewidth=2, color='red', label=f'Rolling Mean (w={{window}})')\n"
            agg_code += f"                axes[i].set_title(f'{{col}} — Trend over {tcol}', fontsize=11)\n"
            agg_code += f"                axes[i].legend()\n"
            agg_code += f"            plt.suptitle('Trend Analysis with Rolling Averages', fontsize=13)\n"
            agg_code += f"            plt.tight_layout()\n"
            agg_code += f"            plt.savefig(os.path.join(r'{output_folder}', 'trend_analysis.png'), dpi=100, bbox_inches='tight')\n"
            agg_code += f"            plt.show(); plt.close()\n"
            agg_code += f"except Exception as e:\n"
            agg_code += f"    print(f'Trend analysis error for {tcol}: {{e}}')\n\n"

    if monotonic:
        agg_code += "# Monotonic trend summary\n"
        agg_code += "print('\\n=== Monotonic Trends (Spearman correlation with row index) ===')\n"
        for col, direction in monotonic.items():
            agg_code += f"print(f'  {col}: {direction} trend')\n"
    elif not time_cols:
        agg_code += "print('No time-series columns or monotonic trends detected.')\n"

    cells.append(new_code_cell(agg_code))

    # ---- Section 12: Pipeline Preparation & Recommendations ----
    cells.append(new_markdown_cell("## Section 12 — Pipeline Preparation & Recommendations"))

    pr = stats.get("pipeline_recommendations", {})
    pipe_code = "# --- Pipeline Configuration Recommendations ---\n"
    pipe_code += "print('=' * 60)\n"
    pipe_code += "print('PIPELINE CONFIGURATION RECOMMENDATIONS')\n"
    pipe_code += "print('=' * 60)\n\n"

    suggested_task = pr.get("suggested_task", "unknown")
    suggested_target = pr.get("suggested_target", "N/A")
    suggested_models = pr.get("suggested_models", [])
    suggested_metric = pr.get("suggested_metric", "N/A")
    suggested_imputer = pr.get("suggested_imputer", "none")
    suggested_scaler = pr.get("suggested_scaler", "StandardScaler")
    suggested_encoder = pr.get("suggested_encoder", "none")
    suggested_transformers = pr.get("suggested_transformers", [])

    pipe_code += f"print('Task Type        : {suggested_task}')\n"
    pipe_code += f"print('Target Column    : {suggested_target}')\n"
    pipe_code += f"print('Recommended Models: {suggested_models}')\n"
    pipe_code += f"print('Evaluation Metric : {suggested_metric}')\n"
    pipe_code += f"print('Imputer Strategy  : {suggested_imputer}')\n"
    pipe_code += f"print('Feature Scaler    : {suggested_scaler}')\n"
    pipe_code += f"print('Categorical Encoder: {suggested_encoder}')\n"
    pipe_code += f"print('Transformers      : {suggested_transformers}')\n"

    if pr.get("class_imbalance"):
        pipe_code += "print('\\n⚠ WARNING: Class imbalance detected!')\n"
        pipe_code += "print('  → Consider SMOTE, class_weight=\"balanced\", or stratified CV.')\n"

    if pr.get("n_classes"):
        n_classes_val = pr["n_classes"]
        pipe_code += f"print('Number of classes : {n_classes_val}')\n"

    pipe_code += "\nprint('\\n' + '=' * 60)\n"
    pipe_code += "print('DATASET READINESS SUMMARY')\n"
    pipe_code += "print('=' * 60)\n"
    pipe_code += "print(f'Final shape       : {cleaned.shape}')\n"
    pipe_code += "print(f'Remaining nulls   : {cleaned.isnull().sum().sum()}')\n"
    pipe_code += "print(f'Remaining dupes   : {cleaned.duplicated().sum()}')\n"
    pipe_code += "print(f'Numeric features  : {len(cleaned.select_dtypes(include=np.number).columns)}')\n"
    pipe_code += "print(f'Categorical feats : {len(cleaned.select_dtypes(include=[\"object\",\"category\"]).columns)}')\n\n"

    pipe_code += "# Feature importance preview (correlation with target if applicable)\n"
    if suggested_target and suggested_target != "N/A" and suggested_task in ("classification", "regression"):
        pipe_code += f"target_col = '{suggested_target}'\n"
        pipe_code += "if target_col in cleaned.columns:\n"
        pipe_code += "    num_feats = cleaned.select_dtypes(include=np.number).columns.tolist()\n"
        pipe_code += "    if target_col in num_feats:\n"
        pipe_code += "        num_feats = [c for c in num_feats if c != target_col]\n"
        pipe_code += "    if num_feats:\n"
        pipe_code += "        correlations = cleaned[num_feats].corrwith(cleaned[target_col].astype(float, errors='ignore') if cleaned[target_col].dtype != object else pd.factorize(cleaned[target_col])[0]).abs().sort_values(ascending=False)\n"
        pipe_code += "        print('\\nFeature-Target Correlations:')\n"
        pipe_code += "        for feat, corr_val in correlations.head(10).items():\n"
        pipe_code += "            bar = '█' * int(corr_val * 20)\n"
        pipe_code += "            print(f'  {feat:30s} {corr_val:.4f} {bar}')\n"

    # Save pipeline config as JSON
    pipe_code += "\n# Save pipeline configuration\n"
    pipe_code += "import json\n"
    pipe_code += "pipeline_config = {\n"
    pipe_code += f"    'task_type': '{suggested_task}',\n"
    pipe_code += f"    'target_column': '{suggested_target}',\n"
    pipe_code += f"    'suggested_models': {repr(suggested_models)},\n"
    pipe_code += f"    'evaluation_metric': '{suggested_metric}',\n"
    pipe_code += f"    'imputer': '{suggested_imputer}',\n"
    pipe_code += f"    'scaler': '{suggested_scaler}',\n"
    pipe_code += f"    'encoder': '{suggested_encoder}',\n"
    pipe_code += f"    'transformers': {repr(suggested_transformers)},\n"
    pipe_code += "    'dataset_shape': list(cleaned.shape),\n"
    pipe_code += "    'feature_columns': cleaned.columns.tolist(),\n"
    pipe_code += "}\n"
    pipe_code += f"config_path = os.path.join(r'{output_folder}', 'pipeline_config.json')\n"
    pipe_code += "with open(config_path, 'w') as f:\n"
    pipe_code += "    json.dump(pipeline_config, f, indent=2, default=str)\n"
    pipe_code += "print(f'\\nPipeline config saved to: {config_path}')\n"

    cells.append(new_code_cell(pipe_code))

    # ---- Section 13: Save Cleaned & Engineered Dataset ----
    cells.append(new_markdown_cell("## Section 13 — Save Pipeline-Ready Dataset"))
    cells.append(new_code_cell(
        f"cleaned_path = r'{output_folder}/cleaned_data.csv'\n"
        "cleaned.to_csv(cleaned_path, index=False)\n"
        "print(f'Pipeline-ready dataset saved to: {cleaned_path}')\n"
        f"print(f'File size: {{os.path.getsize(cleaned_path) / 1024:.1f}} KB')\n"
        "print(f'Final columns ({len(cleaned.columns)}): {cleaned.columns.tolist()}')"
    ))

    nb = new_notebook(cells=cells)
    nb.metadata["kernelspec"] = {
        "display_name": "Python 3",
        "language": "python",
        "name": "python3",
    }
    nb.metadata["language_info"] = {"name": "python", "version": "3.10.0"}

    with open(notebook_path, "w", encoding="utf-8") as f:
        nbformat.write(nb, f)
    logger.info("Notebook written to %s", notebook_path)


# ---------------------------------------------------------------------------
# Notebook execution
# ---------------------------------------------------------------------------

def _execute_notebook(notebook_path: str) -> None:
    """Execute the notebook in-place using jupyter nbconvert."""
    import sys
    # Use the same Python interpreter's jupyter (works in venv)
    jupyter_cmd = [sys.executable, "-m", "jupyter", "nbconvert"]
    try:
        result = subprocess.run(
            jupyter_cmd + [
                "--to", "notebook",
                "--execute",
                "--inplace",
                "--ExecutePreprocessor.timeout=600",
                "--ExecutePreprocessor.kernel_name=python3",
                notebook_path,
            ],
            capture_output=True,
            text=True,
            timeout=620,
        )
        if result.returncode != 0:
            logger.warning("nbconvert stderr: %s", result.stderr[-2000:])
        else:
            logger.info("Notebook executed successfully: %s", notebook_path)
    except FileNotFoundError:
        logger.warning("jupyter nbconvert not found — notebook will not be executed")
    except subprocess.TimeoutExpired:
        logger.warning("Notebook execution timed out for %s", notebook_path)
    except Exception as exc:
        logger.error("Notebook execution error: %s", exc)


# ---------------------------------------------------------------------------
# Word document generation
# ---------------------------------------------------------------------------

def _create_word_doc(
    df: pd.DataFrame,
    stats: Dict[str, Any],
    findings: List[str],
    output_path: str,
) -> None:
    """Generate a Word document summarising the EDA findings."""
    doc = Document()

    # Title
    title = doc.add_heading("EDA Report", 0)
    title.alignment = 1  # center

    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    # Dataset overview
    doc.add_heading("1. Dataset Overview", level=1)
    rows, cols = stats["shape"]
    doc.add_paragraph(f"Shape: {rows:,} rows × {cols} columns")
    doc.add_paragraph(f"Numeric columns: {len(stats['num_cols'])}")
    doc.add_paragraph(f"Categorical columns: {len(stats['cat_cols'])}")
    doc.add_paragraph(f"Duplicate rows: {stats['duplicates']:,}")

    # Data quality table
    doc.add_heading("2. Data Quality — Missing Values", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Column", "Missing Count", "Missing %"
    for col in df.columns:
        cnt = stats["null_counts"].get(col, 0)
        pct = stats["null_pct"].get(col, 0.0)
        row = table.add_row().cells
        row[0].text = col
        row[1].text = str(cnt)
        row[2].text = f"{pct}%"

    # Key findings
    doc.add_heading("3. Key Findings", level=1)
    for f in findings:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f)

    # Top correlations
    if stats["top_correlations"]:
        doc.add_heading("4. Top Correlations", level=1)
        for c1, c2, corr in stats["top_correlations"]:
            doc.add_paragraph(f"• {c1} ↔ {c2}: r = {corr}", style="List Bullet")

    # Skewness
    if stats["skewness"]:
        doc.add_heading("5. Feature Skewness", level=1)
        for col, skew_val in sorted(stats["skewness"].items(), key=lambda x: abs(x[1]), reverse=True)[:10]:
            doc.add_paragraph(f"• {col}: {skew_val:+.4f}")

    # Recommendations
    doc.add_heading("6. Recommendations", level=1)
    recommendations = [
        "Handle missing values appropriately (median/mode imputation or model-based imputation).",
        "Investigate and treat outliers before modelling.",
        "Apply log-transformation to highly skewed numeric features.",
        "Encode categorical variables (one-hot or label encoding) before training.",
        "Use the cleaned_data.csv file for downstream modelling.",
    ]
    for rec in recommendations:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rec)

    # Feature Engineering
    fe = stats.get("feature_engineering", {})
    doc.add_heading("7. Feature Engineering", level=1)
    if fe.get("interaction_candidates"):
        doc.add_paragraph("Interaction features (correlated pairs):")
        for c1, c2 in fe["interaction_candidates"][:5]:
            doc.add_paragraph(f"  • {c1} × {c2}", style="List Bullet")
    if fe.get("log_candidates"):
        doc.add_paragraph(f"Log-transform candidates: {', '.join(fe['log_candidates'][:5])}")
    if fe.get("date_decomp_cols"):
        doc.add_paragraph(f"Date decomposition columns: {', '.join(fe['date_decomp_cols'])}")
    if fe.get("binning_candidates"):
        doc.add_paragraph(f"Binning candidates: {', '.join(fe['binning_candidates'][:5])}")
    if not any(fe.get(k) for k in ["interaction_candidates", "log_candidates", "date_decomp_cols", "binning_candidates"]):
        doc.add_paragraph("No automatic feature engineering opportunities detected.")

    # Aggregation Analysis
    agg = stats.get("aggregation_analysis", {})
    doc.add_heading("8. Aggregation & GroupBy Analysis", level=1)
    if agg.get("groupby_cols"):
        doc.add_paragraph(f"Suitable groupby columns: {', '.join(agg['groupby_cols'])}")
        for grp_col, summary in agg.get("groupby_summaries", {}).items():
            doc.add_paragraph(f"  • {grp_col}: {summary.get('n_groups', '?')} groups, "
                            f"aggregated over {summary.get('agg_columns', [])}")
    else:
        doc.add_paragraph("No categorical columns with 2–20 unique values found for aggregation.")

    # Trend Analysis
    trend = stats.get("trend_analysis", {})
    doc.add_heading("9. Trend Analysis", level=1)
    if trend.get("time_columns"):
        doc.add_paragraph(f"Time-series columns: {', '.join(trend['time_columns'])}")
    if trend.get("monotonic_trends"):
        doc.add_paragraph("Monotonic trends detected:")
        for col, direction in trend["monotonic_trends"].items():
            doc.add_paragraph(f"  • {col}: {direction}", style="List Bullet")
    if not trend.get("time_columns") and not trend.get("monotonic_trends"):
        doc.add_paragraph("No significant trends detected.")

    # Pipeline Recommendations
    pr = stats.get("pipeline_recommendations", {})
    doc.add_heading("10. Pipeline Configuration Recommendations", level=1)
    if pr:
        recs_table = doc.add_table(rows=1, cols=2)
        recs_table.style = "Table Grid"
        hdr = recs_table.rows[0].cells
        hdr[0].text, hdr[1].text = "Parameter", "Recommendation"

        params = [
            ("Task Type", pr.get("suggested_task", "N/A")),
            ("Target Column", pr.get("suggested_target", "N/A")),
            ("Models", ", ".join(pr.get("suggested_models", []))),
            ("Metric", pr.get("suggested_metric", "N/A")),
            ("Imputer", pr.get("suggested_imputer", "N/A")),
            ("Scaler", pr.get("suggested_scaler", "N/A")),
            ("Encoder", pr.get("suggested_encoder", "N/A")),
            ("Transformers", ", ".join(pr.get("suggested_transformers", []))),
        ]
        for param, value in params:
            row = recs_table.add_row().cells
            row[0].text = param
            row[1].text = str(value)

        if pr.get("class_imbalance"):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("⚠ Class imbalance detected — use SMOTE or class_weight='balanced'.").bold = True
    else:
        doc.add_paragraph("Insufficient data to generate pipeline recommendations.")

    doc.save(output_path)
    logger.info("Word document saved to %s", output_path)


# ---------------------------------------------------------------------------
# Data cleaning
# ---------------------------------------------------------------------------

def _clean_dataframe(df: pd.DataFrame, stats: Optional[Dict[str, Any]] = None) -> pd.DataFrame:
    """Return a cleaned copy of *df* with optional feature engineering applied."""
    cleaned = df.copy()
    cleaned = cleaned.drop_duplicates()

    num_cols = cleaned.select_dtypes(include=[np.number]).columns.tolist()
    cat_cols = cleaned.select_dtypes(include=["object", "category"]).columns.tolist()

    for col in num_cols:
        if cleaned[col].isnull().any():
            cleaned[col] = cleaned[col].fillna(cleaned[col].median())
        q1, q3 = cleaned[col].quantile(0.25), cleaned[col].quantile(0.75)
        iqr = q3 - q1
        cleaned[col] = cleaned[col].clip(q1 - 1.5 * iqr, q3 + 1.5 * iqr)

    for col in cat_cols:
        if cleaned[col].isnull().any():
            mode = cleaned[col].mode()
            if not mode.empty:
                cleaned[col] = cleaned[col].fillna(mode[0])

    # Apply feature engineering if stats are provided
    if stats:
        fe = stats.get("feature_engineering", {})

        # Interaction features
        for c1, c2 in fe.get("interaction_candidates", [])[:5]:
            if c1 in cleaned.columns and c2 in cleaned.columns:
                try:
                    cleaned[f"{c1}_x_{c2}"] = cleaned[c1] * cleaned[c2]
                except Exception:
                    pass

        # Log-transform candidates
        for col in fe.get("log_candidates", [])[:5]:
            if col in cleaned.columns:
                try:
                    if (cleaned[col] > 0).all():
                        cleaned[f"{col}_log"] = np.log1p(cleaned[col])
                except Exception:
                    pass

        # Date decomposition
        for col in fe.get("date_decomp_cols", [])[:3]:
            if col in cleaned.columns:
                try:
                    dt = pd.to_datetime(cleaned[col], errors="coerce")
                    if dt.notna().sum() > 0:
                        cleaned[f"{col}_year"] = dt.dt.year
                        cleaned[f"{col}_month"] = dt.dt.month
                        cleaned[f"{col}_dayofweek"] = dt.dt.dayofweek
                        cleaned[f"{col}_quarter"] = dt.dt.quarter
                except Exception:
                    pass

        # Binning candidates
        for col in fe.get("binning_candidates", [])[:5]:
            if col in cleaned.columns:
                try:
                    cleaned[f"{col}_binned"] = pd.qcut(
                        cleaned[col], q=5,
                        labels=["very_low", "low", "mid", "high", "very_high"],
                        duplicates="drop",
                    )
                except Exception:
                    pass

    return cleaned


# ---------------------------------------------------------------------------
# Pipeline config helper
# ---------------------------------------------------------------------------

def _save_pipeline_config(stats: Dict[str, Any], cleaned_df: pd.DataFrame, output_path: str) -> None:
    """Save recommended pipeline configuration as JSON."""
    import json
    pr = stats.get("pipeline_recommendations", {})
    config = {
        "task_type": pr.get("suggested_task", "unknown"),
        "target_column": pr.get("suggested_target", None),
        "suggested_models": pr.get("suggested_models", []),
        "evaluation_metric": pr.get("suggested_metric", "accuracy"),
        "imputer": pr.get("suggested_imputer", "none"),
        "scaler": pr.get("suggested_scaler", "StandardScaler"),
        "encoder": pr.get("suggested_encoder", "none"),
        "transformers": pr.get("suggested_transformers", []),
        "class_imbalance": pr.get("class_imbalance", False),
        "dataset_shape": list(cleaned_df.shape),
        "feature_columns": cleaned_df.columns.tolist(),
        "numeric_columns": cleaned_df.select_dtypes(include=[np.number]).columns.tolist(),
        "categorical_columns": cleaned_df.select_dtypes(include=["object", "category"]).columns.tolist(),
        "feature_engineering_applied": {
            "interactions": [f"{c1}_x_{c2}" for c1, c2 in stats.get("feature_engineering", {}).get("interaction_candidates", [])[:5]],
            "log_transforms": [f"{c}_log" for c in stats.get("feature_engineering", {}).get("log_candidates", [])[:5]],
            "date_features": stats.get("feature_engineering", {}).get("date_decomp_cols", []),
            "binned_features": [f"{c}_binned" for c in stats.get("feature_engineering", {}).get("binning_candidates", [])[:5]],
        },
        "trend_analysis": {
            "time_columns": stats.get("trend_analysis", {}).get("time_columns", []),
            "monotonic_trends": stats.get("trend_analysis", {}).get("monotonic_trends", {}),
        },
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(config, f, indent=2, default=str)
    logger.info("Pipeline config saved to %s", output_path)


# ---------------------------------------------------------------------------
# Zip helper
# ---------------------------------------------------------------------------

def _zip_artifacts(base_folder: str, files: Dict[str, str]) -> str:
    """Zip selected files into base_folder + '.zip' using provided arc names (flattened)."""
    zip_path = base_folder.rstrip("/").rstrip("\\") + ".zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for arcname, src in files.items():
            if src and os.path.exists(src):
                zf.write(src, os.path.basename(arcname))
    logger.info("Artifacts zipped to %s", zip_path)
    return zip_path
