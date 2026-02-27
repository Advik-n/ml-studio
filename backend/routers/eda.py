"""EDA router."""
import logging
import os
from datetime import datetime
from typing import List

import aiofiles
import numpy as np
import pandas as pd
from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, UploadFile, status
from fastapi.responses import FileResponse, JSONResponse
from sqlalchemy.orm import Session

from config import settings
from database import get_db
from models.eda_job import EDAJob
from models.project import Project
from models.user import User
from schemas.eda import EDAJobResponse
from services.eda_service import generate_eda
from utils.dependencies import require_verified_user

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/eda", tags=["EDA"])

ALLOWED_EXTENSIONS = {".csv", ".tsv", ".xls", ".xlsx", ".json", ".parquet"}


async def _run_eda_job(job_id: str, file_path: str, project_folder: str, db_url: str) -> None:
    """Background task: execute EDA and update job record in DB."""
    # Each background task needs its own DB session to avoid threading issues
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker

    engine = create_engine(db_url, connect_args={"check_same_thread": False})
    Session_ = sessionmaker(bind=engine)
    db = Session_()

    job = db.query(EDAJob).filter(EDAJob.id == job_id).first()
    if not job:
        db.close()
        return

    job.status = "processing"
    db.commit()

    try:
        result = await generate_eda(file_path, project_folder, job_id)
        job.status = "completed"
        job.output_folder = result["output_folder"]
        job.notebook_path = result["notebook_path"]
        job.docx_path = result["docx_path"]
        job.cleaned_csv_path = result["cleaned_csv_path"]
        job.zip_path = result["zip_path"]
        job.completed_at = datetime.utcnow()
    except Exception as exc:
        logger.exception("EDA job %s failed: %s", job_id, exc)
        job.status = "failed"
        job.error_message = str(exc)
        job.completed_at = datetime.utcnow()
    finally:
        db.commit()
        db.close()


@router.post("/{project_id}/upload", response_model=EDAJobResponse, status_code=status.HTTP_202_ACCEPTED)
async def upload_for_eda(
    project_id: str,
    file: UploadFile,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(require_verified_user),
    db: Session = Depends(get_db),
):
    """Upload a dataset file and kick off EDA generation in the background."""
    project = db.query(Project).filter(
        Project.id == project_id, Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found.")

    ext = os.path.splitext(file.filename or "")[-1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type '{ext}'. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Persist uploaded file
    uploads_dir = os.path.join(project.folder_path, "uploads")
    os.makedirs(uploads_dir, exist_ok=True)
    dest_path = os.path.join(uploads_dir, file.filename)
    async with aiofiles.open(dest_path, "wb") as out_file:
        while chunk := await file.read(1024 * 1024):  # 1 MB chunks
            await out_file.write(chunk)

    job = EDAJob(
        project_id=project_id,
        input_filename=file.filename,
        status="pending",
    )
    db.add(job)
    db.commit()
    db.refresh(job)

    background_tasks.add_task(
        _run_eda_job, job.id, dest_path, project.folder_path, settings.DATABASE_URL
    )
    return job


@router.get("/{project_id}/jobs", response_model=List[EDAJobResponse])
def list_eda_jobs(
    project_id: str,
    current_user: User = Depends(require_verified_user),
    db: Session = Depends(get_db),
):
    """List all EDA jobs for a project."""
    project = db.query(Project).filter(
        Project.id == project_id, Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found.")
    return db.query(EDAJob).filter(EDAJob.project_id == project_id).all()


@router.get("/jobs/{job_id}", response_model=EDAJobResponse)
def get_eda_job(
    job_id: str,
    current_user: User = Depends(require_verified_user),
    db: Session = Depends(get_db),
):
    """Get the status/result of a single EDA job."""
    job = db.query(EDAJob).filter(EDAJob.id == job_id).first()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="EDA job not found.")
    # Ensure the job belongs to a project owned by the current user
    project = db.query(Project).filter(
        Project.id == job.project_id, Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied.")
    return job


@router.get("/jobs/{job_id}/download")
def download_eda_output(
    job_id: str,
    current_user: User = Depends(require_verified_user),
    db: Session = Depends(get_db),
):
    """Download the zipped EDA output (or notebook if zip unavailable)."""
    job = db.query(EDAJob).filter(EDAJob.id == job_id).first()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="EDA job not found.")
    project = db.query(Project).filter(
        Project.id == job.project_id, Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied.")
    if job.status != "completed":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Job not completed yet.")

    # Prefer zip, fall back to notebook
    download_path = job.zip_path or job.notebook_path
    if not download_path or not os.path.exists(download_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Output file not found on disk.")

    media_type = "application/zip" if download_path.endswith(".zip") else "application/octet-stream"
    return FileResponse(download_path, media_type=media_type, filename=os.path.basename(download_path))


@router.get("/jobs/{job_id}/files/{kind}")
def download_eda_file(
    job_id: str,
    kind: str,
    current_user: User = Depends(require_verified_user),
    db: Session = Depends(get_db),
):
    """Download a specific EDA artifact."""
    job = db.query(EDAJob).filter(EDAJob.id == job_id).first()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="EDA job not found.")
    project = db.query(Project).filter(
        Project.id == job.project_id, Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied.")
    if job.status != "completed":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Job not completed yet.")

    path_map = {
        "zip": job.zip_path or job.notebook_path,
        "docx": job.docx_path,
        "cleaned": job.cleaned_csv_path,
        "notebook": job.notebook_path,
    }
    download_path = path_map.get(kind)
    if not download_path or not os.path.exists(download_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Requested file not found.")

    media_types = {
        "zip": "application/zip",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "cleaned": "text/csv",
        "notebook": "application/octet-stream",
    }
    return FileResponse(
        download_path,
        media_type=media_types.get(kind, "application/octet-stream"),
        filename=os.path.basename(download_path),
    )


def _load_eda_dataframe(job: EDAJob) -> pd.DataFrame:
    """Load the cleaned CSV for an EDA job, falling back to the original upload."""
    csv_path = job.cleaned_csv_path
    if not csv_path or not os.path.exists(csv_path):
        # Fall back to original uploaded file
        project_folder = os.path.dirname(os.path.dirname(csv_path or ""))
        upload_path = os.path.join(project_folder, "uploads", job.input_filename) if project_folder else None
        if upload_path and os.path.exists(upload_path):
            csv_path = upload_path
        else:
            return None
    return pd.read_csv(csv_path)


def _classify_dtype(series: pd.Series) -> str:
    """Classify a pandas Series into one of: numeric, categorical, datetime, text."""
    if pd.api.types.is_numeric_dtype(series):
        return "numeric"
    if pd.api.types.is_datetime64_any_dtype(series):
        return "datetime"
    # Try to detect datetime stored as strings
    if series.dtype == object:
        sample = series.dropna().head(20)
        if len(sample) > 0:
            try:
                pd.to_datetime(sample)
                return "datetime"
            except (ValueError, TypeError):
                pass
        # Text vs categorical heuristic: if many unique values relative to count, it's text
        nunique = series.nunique()
        if nunique > 50 and nunique / max(len(series), 1) > 0.5:
            return "text"
    return "categorical"


def _safe_float(val) -> float:
    """Convert a value to a JSON-safe float."""
    if val is None or (isinstance(val, float) and (np.isnan(val) or np.isinf(val))):
        return 0.0
    return round(float(val), 4)


@router.get("/jobs/{job_id}/data-summary")
def get_data_summary(
    job_id: str,
    current_user: User = Depends(require_verified_user),
    db: Session = Depends(get_db),
):
    """Return a comprehensive data summary for the EDA job's dataset."""
    job = db.query(EDAJob).filter(EDAJob.id == job_id).first()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="EDA job not found.")
    project = db.query(Project).filter(
        Project.id == job.project_id, Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied.")
    if job.status != "completed":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Job not completed yet.")

    df = _load_eda_dataframe(job)
    if df is None or df.empty:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Dataset file not found or empty.")

    # Dataset info
    dataset_info = {
        "rows": len(df),
        "columns": len(df.columns),
        "memory_mb": round(df.memory_usage(deep=True).sum() / (1024 * 1024), 2),
        "duplicates": int(df.duplicated().sum()),
    }

    # Column details
    column_details = []
    col_types = {}
    for col in df.columns:
        dtype = _classify_dtype(df[col])
        col_types[col] = dtype
        missing_count = int(df[col].isna().sum())
        sample_vals = df[col].dropna().unique()[:5].tolist()
        # Make sample values JSON-serializable
        sample_vals = [str(v) if not isinstance(v, (int, float, str, bool)) else v for v in sample_vals]
        column_details.append({
            "name": col,
            "dtype": dtype,
            "missing_count": missing_count,
            "missing_pct": round(missing_count / max(len(df), 1) * 100, 2),
            "unique_count": int(df[col].nunique()),
            "sample_values": sample_vals,
        })

    # Numeric stats
    numeric_cols = [c for c, t in col_types.items() if t == "numeric"]
    numeric_stats = {}
    for col in numeric_cols:
        s = df[col].dropna()
        if len(s) == 0:
            continue
        q1 = _safe_float(s.quantile(0.25))
        q3 = _safe_float(s.quantile(0.75))
        iqr = q3 - q1
        lower = q1 - 1.5 * iqr
        upper = q3 + 1.5 * iqr
        outlier_count = int(((s < lower) | (s > upper)).sum())
        numeric_stats[col] = {
            "mean": _safe_float(s.mean()),
            "std": _safe_float(s.std()),
            "min": _safe_float(s.min()),
            "max": _safe_float(s.max()),
            "median": _safe_float(s.median()),
            "skewness": _safe_float(s.skew()),
            "kurtosis": _safe_float(s.kurtosis()),
            "q1": q1,
            "q3": q3,
            "outlier_count": outlier_count,
        }

    # Categorical stats
    cat_cols = [c for c, t in col_types.items() if t == "categorical"]
    categorical_stats = {}
    for col in cat_cols:
        vc = df[col].value_counts()
        categorical_stats[col] = {
            "unique_count": int(df[col].nunique()),
            "top_value": str(vc.index[0]) if len(vc) > 0 else None,
            "top_frequency": int(vc.iloc[0]) if len(vc) > 0 else 0,
            "value_counts": {str(k): int(v) for k, v in vc.head(10).items()},
        }

    # Correlations (top 10 by abs value)
    correlations = []
    if len(numeric_cols) >= 2:
        corr_matrix = df[numeric_cols].corr()
        pairs = []
        for i, c1 in enumerate(numeric_cols):
            for c2 in numeric_cols[i + 1:]:
                val = corr_matrix.loc[c1, c2]
                if not np.isnan(val):
                    pairs.append({"col1": c1, "col2": c2, "value": _safe_float(val)})
        pairs.sort(key=lambda x: abs(x["value"]), reverse=True)
        correlations = pairs[:10]

    # Recommendations
    recommendations = []
    for col, stats in numeric_stats.items():
        if abs(stats["skewness"]) > 2:
            recommendations.append(f"Column '{col}' has high skewness ({stats['skewness']}) — consider log transform")
        if stats["outlier_count"] > len(df) * 0.05:
            recommendations.append(f"Column '{col}' has {stats['outlier_count']} outliers ({round(stats['outlier_count']/len(df)*100,1)}%) — consider capping or removal")
    for detail in column_details:
        if detail["missing_pct"] > 30:
            recommendations.append(f"Column '{detail['name']}' has {detail['missing_pct']}% missing values — consider dropping or imputing")
        if detail["dtype"] == "categorical" and detail["unique_count"] > 50:
            recommendations.append(f"Column '{detail['name']}' has {detail['unique_count']} categories — consider grouping rare values")
    if dataset_info["duplicates"] > 0:
        recommendations.append(f"Dataset has {dataset_info['duplicates']} duplicate rows — consider removing them")

    # Target suggestions
    classification_targets = []
    regression_targets = []
    for col in df.columns:
        nunique = df[col].nunique()
        if col_types[col] == "categorical" and 2 <= nunique <= 20:
            classification_targets.append(col)
        elif col_types[col] == "numeric":
            if 2 <= nunique <= 20:
                classification_targets.append(col)
            if nunique > 10:
                regression_targets.append(col)

    return JSONResponse(content={
        "dataset_info": dataset_info,
        "column_details": column_details,
        "numeric_stats": numeric_stats,
        "categorical_stats": categorical_stats,
        "correlations": correlations,
        "recommendations": recommendations,
        "target_suggestions": {
            "classification": classification_targets,
            "regression": regression_targets,
        },
    })


@router.get("/jobs/{job_id}/pipeline-report")
def download_pipeline_report(
    job_id: str,
    current_user: User = Depends(require_verified_user),
    db: Session = Depends(get_db),
):
    """Generate and download a Pipeline Preparation Report as a Word document."""
    from docx import Document
    from docx.shared import Inches, Pt

    job = db.query(EDAJob).filter(EDAJob.id == job_id).first()
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="EDA job not found.")
    project = db.query(Project).filter(
        Project.id == job.project_id, Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied.")
    if job.status != "completed":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Job not completed yet.")

    df = _load_eda_dataframe(job)
    if df is None or df.empty:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Dataset file not found or empty.")

    # Classify columns
    col_types = {col: _classify_dtype(df[col]) for col in df.columns}
    numeric_cols = [c for c, t in col_types.items() if t == "numeric"]
    cat_cols = [c for c, t in col_types.items() if t == "categorical"]

    doc = Document()
    doc.add_heading("Pipeline Preparation Report", level=0)

    # Section 1: Dataset Overview
    doc.add_heading("1. Dataset Overview", level=1)
    mem_mb = round(df.memory_usage(deep=True).sum() / (1024 * 1024), 2)
    dupes = int(df.duplicated().sum())
    doc.add_paragraph(f"Rows: {len(df)}")
    doc.add_paragraph(f"Columns: {len(df.columns)}")
    doc.add_paragraph(f"Memory: {mem_mb} MB")
    doc.add_paragraph(f"Duplicate rows: {dupes}")

    # Section 2: Column Details
    doc.add_heading("2. Column Details", level=1)
    tbl = doc.add_table(rows=1, cols=4, style="Table Grid")
    for i, h in enumerate(["Column", "Type", "Missing %", "Unique"]):
        tbl.rows[0].cells[i].text = h
    for col in df.columns:
        row = tbl.add_row()
        missing_pct = round(df[col].isna().sum() / max(len(df), 1) * 100, 2)
        row.cells[0].text = col
        row.cells[1].text = col_types[col]
        row.cells[2].text = str(missing_pct)
        row.cells[3].text = str(df[col].nunique())

    # Section 3: Numeric Feature Statistics
    if numeric_cols:
        doc.add_heading("3. Numeric Feature Statistics", level=1)
        tbl = doc.add_table(rows=1, cols=7, style="Table Grid")
        for i, h in enumerate(["Column", "Mean", "Std", "Skewness", "Kurtosis", "Outliers", "Range"]):
            tbl.rows[0].cells[i].text = h
        for col in numeric_cols:
            s = df[col].dropna()
            if len(s) == 0:
                continue
            q1, q3 = s.quantile(0.25), s.quantile(0.75)
            iqr = q3 - q1
            outliers = int(((s < q1 - 1.5 * iqr) | (s > q3 + 1.5 * iqr)).sum())
            row = tbl.add_row()
            row.cells[0].text = col
            row.cells[1].text = str(_safe_float(s.mean()))
            row.cells[2].text = str(_safe_float(s.std()))
            row.cells[3].text = str(_safe_float(s.skew()))
            row.cells[4].text = str(_safe_float(s.kurtosis()))
            row.cells[5].text = str(outliers)
            row.cells[6].text = f"{_safe_float(s.min())} – {_safe_float(s.max())}"
    else:
        doc.add_heading("3. Numeric Feature Statistics", level=1)
        doc.add_paragraph("No numeric columns found.")

    # Section 4: Categorical Feature Summary
    if cat_cols:
        doc.add_heading("4. Categorical Feature Summary", level=1)
        tbl = doc.add_table(rows=1, cols=4, style="Table Grid")
        for i, h in enumerate(["Column", "Unique", "Top Value", "Frequency"]):
            tbl.rows[0].cells[i].text = h
        for col in cat_cols:
            vc = df[col].value_counts()
            row = tbl.add_row()
            row.cells[0].text = col
            row.cells[1].text = str(df[col].nunique())
            row.cells[2].text = str(vc.index[0]) if len(vc) > 0 else "N/A"
            row.cells[3].text = str(int(vc.iloc[0])) if len(vc) > 0 else "0"
    else:
        doc.add_heading("4. Categorical Feature Summary", level=1)
        doc.add_paragraph("No categorical columns found.")

    # Section 5: Top Correlations
    doc.add_heading("5. Top Correlations", level=1)
    if len(numeric_cols) >= 2:
        corr_matrix = df[numeric_cols].corr()
        pairs = []
        for i, c1 in enumerate(numeric_cols):
            for c2 in numeric_cols[i + 1:]:
                val = corr_matrix.loc[c1, c2]
                if not np.isnan(val):
                    pairs.append((c1, c2, round(float(val), 4)))
        pairs.sort(key=lambda x: abs(x[2]), reverse=True)
        tbl = doc.add_table(rows=1, cols=3, style="Table Grid")
        for i, h in enumerate(["Column 1", "Column 2", "Correlation"]):
            tbl.rows[0].cells[i].text = h
        for c1, c2, v in pairs[:10]:
            row = tbl.add_row()
            row.cells[0].text = c1
            row.cells[1].text = c2
            row.cells[2].text = str(v)
    else:
        doc.add_paragraph("Not enough numeric columns for correlation analysis.")

    # Section 6: Recommendations
    doc.add_heading("6. Recommendations for Pipeline", level=1)
    recs = []
    for col in numeric_cols:
        s = df[col].dropna()
        if len(s) == 0:
            continue
        skew = s.skew()
        if abs(skew) > 2:
            recs.append(f"Column '{col}' has high skewness ({_safe_float(skew)}) — consider log transform")
        q1, q3 = s.quantile(0.25), s.quantile(0.75)
        iqr = q3 - q1
        outlier_count = int(((s < q1 - 1.5 * iqr) | (s > q3 + 1.5 * iqr)).sum())
        if outlier_count > len(df) * 0.05:
            recs.append(f"Column '{col}' has {outlier_count} outliers — consider capping or removal")
    for col in df.columns:
        miss_pct = df[col].isna().sum() / max(len(df), 1) * 100
        if miss_pct > 30:
            recs.append(f"Column '{col}' has {round(miss_pct, 1)}% missing values — consider dropping or imputing")
    if not recs:
        recs.append("Dataset looks clean — no major issues detected.")
    for r in recs:
        doc.add_paragraph(r, style="List Bullet")

    # Section 7: Suggested Targets
    doc.add_heading("7. Suggested Targets", level=1)
    cls_targets, reg_targets = [], []
    for col in df.columns:
        nunique = df[col].nunique()
        if col_types[col] == "categorical" and 2 <= nunique <= 20:
            cls_targets.append(col)
        elif col_types[col] == "numeric":
            if 2 <= nunique <= 20:
                cls_targets.append(col)
            if nunique > 10:
                reg_targets.append(col)
    doc.add_paragraph(f"Classification candidates: {', '.join(cls_targets) if cls_targets else 'None identified'}")
    doc.add_paragraph(f"Regression candidates: {', '.join(reg_targets) if reg_targets else 'None identified'}")

    # Save document
    output_folder = job.output_folder or os.path.dirname(job.cleaned_csv_path or "")
    os.makedirs(output_folder, exist_ok=True)
    report_path = os.path.join(output_folder, "pipeline_preparation_report.docx")
    doc.save(report_path)

    return FileResponse(
        report_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="pipeline_preparation_report.docx",
    )
